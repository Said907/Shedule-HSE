import telebot
from telebot import types
import sqlite3
import json
import requests
import datetime
import os
from docx import Document

user_markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
TOKEN = '6879038917:AAEqTCVsNRIhsDVqv6dz3M5j3ePXZwxtXDs'

first_value = 0


def get_id(message):  # ***

    """This function is used to retrieve the group id from a SQLite database based on the group name.

    :param message: The message containing the group name
    :type message: str

    :returns: The group id corresponding to the group name
    :rtype: int """

    conn = sqlite3.connect('hse.db')
    cursor = conn.cursor()
    cursor.execute("SELECT groupid FROM grup WHERE grnm = ?", (message.text,))
    result = cursor.fetchall()
    global first_value
    first_value = result[0][0]


bot = telebot.TeleBot(TOKEN)
schedule = {}
notes = {}
selected_day = None
current_filter = None


@bot.message_handler(commands=['start'])
def handle_start(message):  # ***

    """ This function is used to handle the start of a process.

    :param message: The message to start the process.
    :type message: str """

    ask_course(message)


@bot.message_handler(func=lambda message: True)
def handle_messages(message):
    """ This function handles various messages and performs corresponding actions.

    :param message: The message to be handled.
    :type message: str """

    conn = sqlite3.connect('hse.db')
    cursor = conn.cursor()
    cursor.execute("SELECT grnm FROM grup")
    result = cursor.fetchall()
    cleaned_result = [row[0] for row in result]
    if message.text == 'Расписание':
        show_schedule_keyboard(message)
    elif message.text == 'Вернуться':
        show_main_keyboard(message)
    elif message.text == 'Заметка':
        ask_for_note(message)
    elif message.text == 'Просмотреть заметки':
        show_notes(message)
    elif message.text == 'Выбрать день':
        ask_for_day(message)
    elif message.text == 'Фильтр':
        show_filter_keyboard(message)
    elif message.text == 'Фильтр по преподу':
        show_subject_filter(message)
    elif message.text == 'Фильтр по предмету':
        show_lecturer_filter(message)
    elif message.text == 'Фильтр по дате':
        show_date_filter(message)
    elif message.text == 'Фильтр по типу занятий':
        show_lesson_type_filter(message)
    elif message.text == 'Комбинированный фильтр':
        ask_combined_filter_params(message)
    elif message.text == 'Вернуться к фильтру':
        show_filter_keyboard(message)
    elif message.text == 'Курс 1':
        ask_group(message)
    elif message.text == 'Курс 2':
        ask_group(message)
    elif message.text == 'Курс 3':
        ask_group(message)
    elif message.text == 'Курс 4':
        ask_group(message)
    elif message.text == 'Поиск по ID урока':
        ask_for_lesson_id_filter(message)
    elif message.text in cleaned_result:
        get_id(message)
        show_main_keyboard(message)
    elif message.text in schedule.keys():
        show_subjects(message)
    else:
        bot.send_message(message.from_user.id, 'Я не понимаю, что вы от меня хотите.')
    conn.close()


def ask_course(message):  # ***

    """ Function to ask user to choose their course

    :param message: :class: `types.Message`
    :returns: None """

    user_markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
    button_course = types.KeyboardButton('Курс 1')
    button_course1 = types.KeyboardButton('Курс 2')
    button_course2 = types.KeyboardButton('Курс 3')
    button_course3 = types.KeyboardButton('Курс 4')
    user_markup.add(button_course, button_course1, button_course2, button_course3)
    bot.send_message(message.from_user.id, 'Привет! Выбери свой курс:', reply_markup=user_markup)


def ask_group(message):  # ***

    """ Function to ask user to choose their group based on the selected course

   :param message: :class:`types.Message`
   :return: None """

    conn = sqlite3.connect('hse.db')
    cursor = conn.cursor()
    user_markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
    group_id = 0
    if message.text == 'Курс 1':
        group_id = 1
        cursor.execute("SELECT grnm FROM grup WHERE id = ?", (group_id,))

        result = cursor.fetchall()
        for row in result:
            button = telebot.types.KeyboardButton(text=str(row[0]))
            user_markup.add(button)
        bot.send_message(message.from_user.id, 'Выберите группу:', reply_markup=user_markup)
    if message.text == 'Курс 2':
        group_id = 2
        cursor.execute("SELECT grnm FROM grup WHERE id = ?", (group_id,))

        result = cursor.fetchall()
        for row in result:
            button = telebot.types.KeyboardButton(text=str(row[0]))
            user_markup.add(button)
        bot.send_message(message.from_user.id, 'Выберите группу:', reply_markup=user_markup)
    if message.text == 'Курс 3':
        group_id = 3
        cursor.execute("SELECT grnm FROM grup WHERE id = ?", (group_id,))

        result = cursor.fetchall()
        for row in result:
            button = telebot.types.KeyboardButton(text=str(row[0]))
            user_markup.add(button)
        bot.send_message(message.from_user.id, 'Выберите группу:', reply_markup=user_markup)
    if message.text == 'Курс 4':
        group_id = 4
        cursor.execute("SELECT grnm FROM grup WHERE id = ?", (group_id,))

        result = cursor.fetchall()
        for row in result:
            button = telebot.types.KeyboardButton(text=str(row[0]))
            user_markup.add(button)
        bot.send_message(message.from_user.id, 'Выберите группу:', reply_markup=user_markup)
    conn.close()


def show_main_keyboard(message):  # ***

    """ Function to display the main keyboard with options to the user

    :param message: :class:`types.Message`
    :return: None """

    url = f'https://www.hse.ru/api/timetable/lessons?fromdate=2023.11.01&todate=2023.12.31&groupoid={first_value}&receiverType=3'
    response = requests.get(url)
    if response.status_code == 200:
        global data
        data = json.loads(response.text)
    user_markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
    button_schedule = types.KeyboardButton('Расписание')
    user_markup.add(button_schedule)
    bot.send_message(message.from_user.id, 'Выбери действие', reply_markup=user_markup)


def get_all_lesson_types_date(selected_date):  # ***

    """ This function takes a selected date as input and returns a list of all lesson types for that date.

    :param selecteddate: A string representing the date for which to retrieve lesson types.
    :type selecteddate: str

    :returns: A list of all lesson types for the selected date, excluding lessons with the discipline "Английский".
    :rtype: list

    :raises ValueError: If the input format for selecteddate is not valid. """

    return [lesson for lesson in data['Lessons'] if
            lesson["date"] == selected_date and not "Английский" in lesson["discipline"]]


def get_all_lesson_types_lecturer(selected_lecturer):  # ***

    """ This function takes a selected lecturer as input and returns a list of all lesson types for that lecturer.

   :param selectedlecturer: A string representing the lecturer for which to retrieve lesson types.
   :type selectedlecturer: str

   :returns: A list of all lesson types for the selected lecturer, excluding lessons with the discipline "Английский".
   :rtype: list

   :raises ValueError: If the input format for selectedlecturer is not valid. """

    return [lesson for lesson in data['Lessons'] if
            lesson["lecturer"] == selected_lecturer and not "Английский" in lesson["discipline"]]


def get_all_lesson_types_subject(selected_subject):  # ***

    """ This function takes a selected subject as input and returns a list of all lesson types for that subject.

    :param selected_subject: A string representing the subject for which to retrieve lesson types.
    :type selected_subject: str

    :returns: A list of all lesson types for the selected subject.
    :rtype: list """

    return [lesson for lesson in data['Lessons'] if selected_subject in lesson["discipline"]]


def show_schedule_keyboard(message):  # ***

    """ This function generates a custom keyboard and handles user actions related to selecting options from the keyboard.

    :param message: The incoming message from the user triggering the keyboard display.
    :type message: str """

    user_markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
    button_notes = types.KeyboardButton('Заметка')
    button_filter = types.KeyboardButton('Фильтр')
    button_return = types.KeyboardButton('Вернуться')
    button_show_notes = types.KeyboardButton('Просмотреть заметки')
    user_markup.add(button_notes, button_filter, button_return, button_show_notes)
    if message.text == 'Расписание':
        current_date = datetime.datetime.now()
        selected_date = current_date.strftime("%Y.%m.%d")
        filtered_lessons = get_all_lesson_types_date(selected_date)
        if filtered_lessons:
            send_filtered_lessons(message, filtered_lessons)
        else:
            bot.send_message(message.chat.id, 'На сегодня расписания нет.')

    bot.send_message(message.chat.id, 'Выберите действие:', reply_markup=user_markup)


def create_lesson_id(filtered_lessons):  # ***

    """ This function takes a list of filtered lessons and generates a unique lesson identifier for each lesson based on the lesson's date, auditorium, beginLesson,
    and endLesson.

    :param filtered_lessons:  A list of dictionaries containing lesson information.
    :type filtered_lessons: list
    :returns filtered_lessons: A list of unique lesson identifiers.
    :rtype lesson_id: list """

    lesson_id = [
        f"{lesson['date'].replace('.', '')}{lesson['auditorium'].replace(' ', '  ')}{lesson['beginLesson'].replace(':', '')}{lesson['endLesson'].replace(':', '')}"
        for lesson in filtered_lessons
    ]
    return lesson_id


conn = sqlite3.connect('hse.db')
cursor = conn.cursor()

cursor.execute('''
    CREATE TABLE IF NOT EXISTS note (
        id INTEGER PRIMARY KEY,
        lesson_id TEXT,
        user_id INTEGER,
        note_text TEXT
    )
''')
conn.commit()


def ask_for_note(message):  # ***

    """ This function sends a message to the user asking them to input the lesson ID for adding a note and registers the next step handler to call the `save_note` function.
    :param message:  The message object received from the user.
    :type message: str
    :returns: none
    :raises: This function does not raise any exceptions or errors. """

    bot.send_message(message.from_user.id, 'Введите id урока для записи заметки:')
    bot.register_next_step_handler(message, save_note)


def save_note(message):  # ***

    """ This function handles saving a note to the database.

    :param message: The message object from the user.
    :type message: str

    :returns: None """

    lesson_id = message.text
    if lesson_id.replace(" ", ""):
        bot.send_message(message.from_user.id, 'Введите заметку:')
        bot.register_next_step_handler(message,
                                       lambda msg: save_note_to_database(lesson_id, message.from_user.id, msg.text,
                                                                         message))
    else:
        bot.send_message(message.from_user.id, 'Неверный формат ID. Пожалуйста, введите числовой ID без пробелов.')
        bot.register_next_step_handler(message, save_note)


def save_note_to_database(lesson_id, user_id, note_text, message):  # ***

    """ This function handles saving a note to the database.

    :param message: The message object from the user.
    :type message: str

    :returns: None """

    conn = sqlite3.connect('hse.db')
    cursor = conn.cursor()

    cursor.execute('INSERT INTO note (lesson_id, user_id, note_text) VALUES (?, ?, ?)', (lesson_id, user_id, note_text))

    conn.commit()
    show_schedule_keyboard(message)


def show_notes(message):
    """ This function handles displaying the notes from the database for a specific user.

    :param message: The message object from the user.
    :type message: str

    :returns: None """

    user_id = message.from_user.id
    conn = sqlite3.connect('hse.db')
    cursor = conn.cursor()

    cursor.execute('SELECT lesson_id, note_text FROM note WHERE user_id = ?', (user_id,))
    notes = cursor.fetchall()

    if not notes:
        bot.send_message(message.from_user.id, 'Заметок пока нет.')
    else:
        notes_text = '\n'.join([f'{lesson_id}: {note_text}' for lesson_id, note_text in notes])
        bot.send_message(message.from_user.id, 'Ваши заметки:\n' + notes_text)


def ask_combined_filter_params(message):
    """  Function that prompts the user to enter parameters for a combined filter and registers the next step handler for processing the input.

    :param message: The incoming message from the user.
    :type message: types.Message
    :returns: None """

    user_markup = types.ReplyKeyboardRemove()
    bot.send_message(message.from_user.id,
                     'Введите параметры для комбинированного фильтра (через запятую): Преподаватель,Предмет,Дата,Тип занятий (введите от 2 до 4 параметров)',
                     reply_markup=user_markup)
    bot.register_next_step_handler(message, process_combined_filter_params)


def ask_combined_filter_params(message):
    """ This function removes the reply keyboard and sends a message to the user to enter parameters in the format "Parameter: Value" separated by commas. It then registers the next step handler for processing the input.

   :param message: The incoming message from the user.
   :type message: types.Message
   :returns: None """

    user_markup = types.ReplyKeyboardRemove()
    bot.send_message(message.from_user.id,
                     'Введите параметры в формате "Параметр: Значение" через запятую. (введите от 2 до 4 параметров)',
                     reply_markup=user_markup)
    bot.register_next_step_handler(message, process_combined_filter_params)


def process_combined_filter_params(message):
    """ This function processes combined filter parameters from the user's message.

    :param message: A string representing the user's message.

    :return: None """

    params_dict = {param.split(':')[0].strip(): param.split(':')[1].strip() for param in message.text.split(',')}
    selected_lecturer = params_dict.get('Преподаватель')
    selected_subject = params_dict.get('Предмет')
    selected_date = params_dict.get('Дата')
    selected_lesson_type = params_dict.get('Тип занятия')

    first_word = next(iter(params_dict))

    if first_word == 'Преподаватель':
        filtered_lessons = get_all_lesson_types_lecturer(selected_lecturer)
        if selected_subject:
            filtered_lessons = [lesson for lesson in filtered_lessons if selected_subject in lesson['discipline', '']]
        if selected_date:
            filtered_lessons = [lesson for lesson in filtered_lessons if lesson['date'] == selected_date]
        if selected_lesson_type:
            filtered_lessons = [lesson for lesson in filtered_lessons if lesson['kindOfWork'] == selected_lesson_type]
    elif first_word == 'Предмет':
        filtered_lessons = get_all_lesson_types_subject(selected_subject)
        if selected_lecturer:
            filtered_lessons = [lesson for lesson in filtered_lessons if lesson['lecturer'] == selected_lecturer]
        if selected_date:
            filtered_lessons = [lesson for lesson in filtered_lessons if lesson['date'] == selected_date]
        if selected_lesson_type:
            filtered_lessons = [lesson for lesson in filtered_lessons if lesson['kindOfWork'] == selected_lesson_type]
    elif first_word == 'Дата':
        filtered_lessons = get_all_lesson_types_date(selected_date)
        if selected_subject:
            filtered_lessons = [lesson for lesson in filtered_lessons if selected_subject in lesson['discipline', '']]
        if selected_lecturer:
            filtered_lessons = [lesson for lesson in filtered_lessons if lesson['lecturer'] == selected_lecturer]
        if selected_lesson_type:
            filtered_lessons = [lesson for lesson in filtered_lessons if lesson['kindOfWork'] == selected_lesson_type]
    elif first_word == 'Тип занятий':
        filtered_lessons = filter_lessons_by_type(selected_lesson_type)
        if selected_subject:
            filtered_lessons = [lesson for lesson in filtered_lessons if selected_subject in lesson['discipline', '']]
        if selected_lecturer:
            filtered_lessons = [lesson for lesson in filtered_lessons if lesson['lecturer'] == selected_lecturer]
        if selected_date:
            filtered_lessons = [lesson for lesson in filtered_lessons if lesson['date'] == selected_date]
    else:
        bot.send_message(message.from_user.id,
                         'Первый параметр должен быть "Преподаватель", "Предмет", "Дата" или "Тип занятия"')
        return

    if filtered_lessons:
        send_filtered_lessons(message, filtered_lessons)
        show_filter_keyboard(message)
    else:
        answerempty(message)
        show_filter_keyboard(message)


def ask_for_day(message):  # ***

    """ This function prompts the user to select a day from the schedule.

    :param message: The message object from the user.
    :type message: str

    :returns: None """

    user_markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
    for day in schedule.keys():
        user_markup.add(types.KeyboardButton(day))

    bot.send_message(message.from_user.id, 'Выберите день:', reply_markup=user_markup)
    bot.register_next_step_handler(message, get_selected_day)


def get_selected_day(message):  # ***

    """ This function stores the selected day and proceeds to show the subjects for the selected day.

    :param message: The message object from the user containing the selected day.
    :type message: str

    :returns: None """

    global selected_day
    selected_day = message.text.lower()
    show_subjects(message)


def show_subjects(message):  # ***

    """ This function is used to display subjects in the schedule for a specific day.

    :param message: Telegram message object
    :returns: None """

    if selected_day in schedule:
        subjects = [subject.split(':')[1].strip() for subject in schedule[selected_day].split('\n')]
        user_markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
        for subject in subjects:
            user_markup.add(types.KeyboardButton(subject))

        user_markup.add(types.KeyboardButton('Вернуться'))

        bot.send_message(message.from_user.id, f'Выберите предмет на {selected_day} для просмотра расписания:',
                         reply_markup=user_markup)
    else:
        bot.send_message(message.from_user.id, 'Ошибка выбора дня.')


def show_filter_keyboard(message):
    """ This function is used to display a keyboard with filtering options to the user.

    :param message: Telegram message object

    :returns: None """

    user_markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
    button_subject_filter = types.KeyboardButton('Фильтр по преподу')
    button_lecturer_filter = types.KeyboardButton('Фильтр по предмету')
    button_date_filter = types.KeyboardButton('Фильтр по дате')
    button_lesson_type_filter = types.KeyboardButton('Фильтр по типу занятий')
    button_combined_filter = types.KeyboardButton('Комбинированный фильтр')
    button_id_filter = types.KeyboardButton('Поиск по ID урока')  # Add this option
    button_return = types.KeyboardButton('Вернуться')
    user_markup.add(button_subject_filter, button_lecturer_filter, button_date_filter,
                    button_lesson_type_filter, button_id_filter, button_combined_filter, button_return)

    bot.send_message(message.from_user.id, 'Выберите фильтр:', reply_markup=user_markup)


def ask_for_lesson_id_filter(message):  # ***

    """ This function prompts the user to enter a lesson ID for filtering.

    :param message: Telegram message object

    :returns: None """

    user_markup = types.ReplyKeyboardRemove()
    bot.send_message(message.from_user.id, 'Введите ID урока для фильтрации:', reply_markup=user_markup)
    bot.register_next_step_handler(message, filter_lessons_by_id)


def get_lesson_by_id(lesson_id):  # ***

    """ This function takes a lesson ID as input and returns the lesson with the matching ID from the data.

   :param lesson_id: The ID of the lesson to retrieve.
   :type lesson_id: int or str

   :returns: The lesson with the matching ID, or None if no lesson is found.
   :rtype: dict or None """

    return next((lesson for lesson in data['Lessons'] if create_lesson_id([lesson])[0] == lesson_id), None)


def filter_lessons_by_id(message):  # ***

    """ This function filters lessons by their ID and sends the filtered lesson to the user. If the lesson with the specified ID is not found, it sends a message to the user indicating that the lesson was not found.

    :param message: The message object containing the lesson ID.
    :type message: object

    :returns: None """

    lesson_id = message.text
    filtered_lesson = get_lesson_by_id(lesson_id)
    if filtered_lesson:
        send_filtered_lessons(message, [filtered_lesson])
        show_filter_keyboard(message)
    else:
        bot.send_message(message.from_user.id, f'Урок с ID {lesson_id} не найден.')
        show_filter_keyboard(message)


def show_subject_filter(message):  # ***

    """  This function displays a prompt to the user to enter the name of the lecturer for filtering.

    :param message: The message object.
    :type message: object

    :returns: None """

    markup = types.ReplyKeyboardRemove()
    bot.send_message(message.from_user.id, 'Введите имя преподавателя для фильтрации:', reply_markup=markup)
    bot.register_next_step_handler(message, handle_subject_filter_lectorer)


def handle_subject_filter_lectorer(message):  # ***

    """ This function handles the subject filter for lecturers.

    :param message: The message containing the selected lecturer.
    :type message: str

    :returns: None """

    selected_lecturer = message.text
    filtered_lessons = get_all_lesson_types_lecturer(selected_lecturer)
    if filtered_lessons:
        send_filtered_lessons(message, filtered_lessons)
        show_schedule_keyboard(message)
    else:
        answerempty(message, filtered_lessons)
        show_schedule_keyboard(message)


def show_lecturer_filter(message):  # ***

    """ This function shows the lecturer filter to the user.

    :param message: The message triggering the function.
    :type message: str

    :returns: None """

    markup = types.ReplyKeyboardRemove()
    bot.send_message(message.from_user.id, 'Введите имя предмет для фильтрации:', reply_markup=markup)
    bot.register_next_step_handler(message, handle_subject_filter_subject)


def handle_subject_filter_subject(message):  # ***

    """ This function handles the subject filter for subjects.

    :param message: The message containing the selected subject.
    :type message: str

    :returns: None """

    selected_subject = message.text
    filtered_lessons = get_all_lesson_types_subject(selected_subject)
    if filtered_lessons:
        send_filtered_lessons(message, filtered_lessons)
        show_schedule_keyboard(message)
    else:
        answerempty(message, filtered_lessons)
        show_schedule_keyboard(message)


def show_date_filter(message):  # ***

    """ This function shows the date filter to the user.

    :param message: The message triggering the function.
    :type message: str

    :returns: None """

    user_markup = types.ReplyKeyboardMarkup(resize_keyboard=True, one_time_keyboard=False)
    button_cancel = types.KeyboardButton('Отмена')
    user_markup.add(button_cancel)
    bot.send_message(message.from_user.id, 'Выберите дату для фильтрации:(В формате гггг.мм.дд)',
                     reply_markup=user_markup)
    bot.register_next_step_handler(message, process_date_filter_response)


def process_date_filter_response(message):  # ***

    """  This function processes the user's response to the date filter.

    :param message: The message containing the user's response.
    :type message: str

    :returns: None """

    if message.text.lower() == 'отмена':
        bot.send_message(message.from_user.id, 'Отменено.')
        show_schedule_keyboard(message)
    else:
        selected_date = message.text
        filtered_lessons = get_all_lesson_types_date(selected_date)
        if filtered_lessons:
            send_filtered_lessons(message, filtered_lessons)
            show_schedule_keyboard(message)
        else:
            answerempty(message, filtered_lessons)
            show_schedule_keyboard(message)


def show_lesson_type_filter(message):  # ***

    """  This function creates and displays a keyboard with buttons for different types of lessons (seminar, lecture, practice, test) and a button to return to the main filter menu.
   It also registers the next step handler to process the user's choice of lesson type.

   :param message: The message object that triggered the function.
   :type message: :class: `types.Message` """

    user_markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
    button_seminar = types.KeyboardButton('Семинары')
    button_lecture = types.KeyboardButton('Лекция')
    button_practice = types.KeyboardButton('Практика')
    button_test = types.KeyboardButton('Экзамен')
    button_return = types.KeyboardButton('Вернуться к фильтру')
    user_markup.add(button_seminar, button_lecture, button_practice, button_return, button_test)

    bot.send_message(message.from_user.id, 'Выберите тип занятий для фильтрации:', reply_markup=user_markup)
    bot.register_next_step_handler(message, process_lesson_type_choices)


def process_lesson_type_choices(message):  # ***

    """  This function processes the user's choice of lesson type and takes appropriate actions based on the choice.

   :param message: The message object containing the user's choice of lesson type.
   :type message: :class: types.Message """


    if message.text in ['Семинары', 'Лекция', 'Практика', 'Экзамен']:
        lesson_type = message.text
        filtered_lessons = filter_lessons_by_type(lesson_type)
        send_filtered_lessons(message, filtered_lessons)
        show_filter_keyboard(message)
    elif message.text == 'Вернуться к фильтру':
        show_filter_keyboard(message)
    else:
        bot.send_message(message.from_user.id, 'Выберите тип занятий, используя предоставленные кнопки.')


def filter_lessons_by_type(lesson_type):  # ***

    """ This function filters lessons based on the provided lesson type.

    :param lesson_type: The type of lesson to filter by.
    :type lesson_type: str

    :returns: A list of lessons that match the specified lesson type.
    :rtype: list """

    if lesson_type != 'Практика':
        return [lesson for lesson in data['Lessons'] if lesson['kindOfWork'] == lesson_type]
    else:
        return [lesson for lesson in data['Lessons'] if lesson['kindOfWork'] == 'Практические занятия']


def answerempty(message):  # ***

    """  Sends a message indicating that the user has no classes on the specified day and displays the keyboard to show the schedule.

    :param message: object representing the user's message
    :type message: str """

    bot.send_message(message.from_user.id, "В этот день у вас занятий нет")
    show_schedule_keyboard(message)


def send_filtered_lessons(message, filtered_lessons):  # ***

    """ Sends filtered lessons data based on the specified filter type to the user.

    :param message: user message object
    :type message: object

    :param filteredlessons: list of filtered lesson objects
    :type filteredlessons: list """

    try:
        if filtered_lessons:
            id_generator = create_lesson_id(filtered_lessons)
            result_text = ""

            for lesson, lesson_id in zip(filtered_lessons, id_generator):
                result_text += (
                    f"Группа: {lesson['group']}\n"
                    f"Предмет: {lesson['discipline']}\n"
                    f"ID Предмета: `{lesson_id}`\n"
                    f"Тип занятия: {lesson['kindOfWork']}\n"
                    f"Преподаватель: {lesson['lecturer']}\n"
                    f"Дата: {lesson['date']}\n"
                    f"Аудитория: {lesson['auditorium']}\n"
                    f"Время: {lesson['beginLesson']} - {lesson['endLesson']}\n"
                    f"-----------------------------------------------------\n"
                )

            bot.send_message(message.chat.id, f'Отфильтрованные данные по типу занятий "{message.text}":\n'
                                              f'-----------------------------------------------------\n'
                                              f'{result_text}', parse_mode='Markdown')
        else:
            bot.send_message(message.chat.id, f'Нет данных по "{message.text}".')
            show_schedule_keyboard(message)

    except:
        document = Document()
        document.add_heading(f'Отфильтрованные данные по "{message.text}"', 0)

        for lesson, lesson_id in zip(filtered_lessons, create_lesson_id(filtered_lessons)):
            document.add_paragraph(f"Группа: {lesson['group']}")
            document.add_paragraph(f"Предмет: {lesson['discipline']}")
            document.add_paragraph(f"ID Предмета: {lesson_id}")
            document.add_paragraph(f"Тип занятия: {lesson['kindOfWork']}")
            document.add_paragraph(f"Преподаватель: {lesson['lecturer']}")
            document.add_paragraph(f"Дата: {lesson['date']}")
            document.add_paragraph(f"Аудитория: {lesson['auditorium']}")
            document.add_paragraph(f"Время: {lesson['beginLesson']} - {lesson['endLesson']}")
            document.add_paragraph('\n')

        file_path = f'filtered_lessons_{message.from_user.id}.docx'
        document.save(file_path)

        with open(file_path, 'rb') as docx_file:
            bot.send_document(message.chat.id, docx_file)
            os.remove(file_path)


if __name__ == '__main__':
    bot.polling(none_stop=True)
